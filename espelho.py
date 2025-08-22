from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.section import WD_ORIENT
import tempfile
import os

import win32com.client


# Função gerar_espelho sem o parâmetro 'data'
def gerar_espelho(referencia, canal, semana, marca):
    # Cria documento
    doc = Document()

    # Página em modo paisagem
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Pt(36)

    # Título com fundo cinza
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run("📦 RECEBIMENTO CD PAVUNA")
    title_run.bold = True
    title_run.font.size = Pt(46)
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    title_paragraph._p.get_or_add_pPr().append(shading_elm)

    # Espaço
    doc.add_paragraph()

    # Tabela para os dados
    table = doc.add_table(rows=3, cols=2)
    table.autofit = False
    table.columns[0].width = Pt(300)
    table.columns[1].width = Pt(300)

    dados = [
        ("📌 REFERÊNCIA:", referencia),
        ("👗 MARCA:", marca),
        ("📈 SEMANA LOJA:", semana),
    ]

    for i, (label, valor) in enumerate(dados):
        if label == "📈 SEMANA LOJA:":
            cell1 = table.cell(i, 0)
            cell2 = table.cell(i, 1)
            cell1.merge(cell2)

            merged_paragraph = cell1.paragraphs[0]
            merged_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            run_label = merged_paragraph.add_run(label)
            run_label.bold = True
            run_label.font.size = Pt(45)
            run_label.font.color.rgb = RGBColor(0, 0, 0)

            run_label.add_break()

            run_value = merged_paragraph.add_run(valor)
            run_value.bold = True
            run_value.font.size = Pt(90)
            run_value.font.color.rgb = RGBColor(0, 0, 0)

            merged_paragraph.paragraph_format.space_before = Pt(60)

        else:
            cell_label = table.cell(i, 0).paragraphs[0]
            run_label = cell_label.add_run(label)
            run_label.bold = True
            run_label.font.size = Pt(40)
            run_label.font.color.rgb = RGBColor(0, 0, 0)

            cell_value = table.cell(i, 1).paragraphs[0]
            run_value = cell_value.add_run(valor)
            run_value.bold = True
            run_value.font.size = Pt(50)
            run_value.font.color.rgb = RGBColor(0, 0, 0)

    # Salva temporariamente
    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    tmp_path = tmp_file.name
    tmp_file.close()
    doc.save(tmp_path)

    return tmp_path


def imprimir_espelho(tmp_path):
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(tmp_path)
    doc.PrintOut()
    doc.Close(False)
    word.Quit()
    os.remove(tmp_path)


if __name__ == "__main__":
    # Preencha aqui os valores ou chame a partir de sua interface
    referencia = "123458907"
    # data = "10/07/2025" # REMOVIDO
    canal = "VAREJO"
    semana = "4"
    marca = "FARM"

    tmp_doc = gerar_espelho(referencia, canal, semana, marca) # 'data' removida
    imprimir_espelho(tmp_doc)

    print("✅ Espelho gerado e enviado para impressão.")